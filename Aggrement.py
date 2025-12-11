import pandas as pd
from docx import Document
import tkinter as tk
from tkinter import filedialog
import os
import sys
from datetime import datetime
import shutil # File operations ke liye
import win32com.client as win32 
# -----------------------------------------------------------
FILENAME_COLUMN = 'Merchant_Code'
SHEET_NAME = 'Sheet1' 
wdExportFormatPDF = 17 

# -----------------------------------------------------------
# 2. Filename safai ka function
def clean_filename(filename):
    """Invalid characters."""
    invalid_chars = '/\\:*?"<>|'
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    return filename.strip()

# -----------------------------------------------------------
# 3. File selection dialog function
def select_file(title, filetypes):
    """File dialog open."""
    root = tk.Tk()
    root.withdraw() 
    
    file_path = filedialog.askopenfilename(
        title=title,
        filetypes=filetypes
    )
    
    if not file_path:
        print(f"Error: its Done.")
        sys.exit()

    return file_path

# -----------------------------------------------------------
# 4. Text replacement function
def replace_text_in_all_elements(document, placeholder, replacement):
    """Document Story"""
    
    
    story_paragraphs = []
    
    
    story_paragraphs.extend(document.paragraphs)
    
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                story_paragraphs.extend(cell.paragraphs)

    # 3. Headers aur Footers ke paragraphs
    for section in document.sections:
        # Headers
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                story_paragraphs.extend(header.paragraphs)
        # Footers
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                story_paragraphs.extend(footer.paragraphs)

    
    for paragraph in story_paragraphs:
        # Run-level replacement 
        for run in paragraph.runs:
            run.text = run.text.replace(placeholder, replacement)


# -----------------------------------------------------------
# Main execution block
if __name__ == "__main__":
    
    # A. 
    SCRIPT_DIR = os.path.dirname(os.path.abspath(sys.argv[0]))
    OUTPUT_FOLDER = SCRIPT_DIR + os.sep + "Output_MailMerge" + os.sep
    word_app = None

    # B. Files
    print("Step 1/2: Word Template file chunen...")
    TEMPLATE_FILE_PATH = select_file(
        title="1. Kripya Word Mail Merge Template file (.docx) chunen",
        filetypes=(("Word files", "*.docx"), ("All files", "*.*"))
    )

    print("Step 2/2: Excel Data file chunen...")
    DATA_FILE_PATH = select_file(
        title="2. Kripya Excel Data file (.xlsx) chunen",
        filetypes=(("Excel files", "*.xlsx"), ("All files", "*.*"))
    )

    # C. Output folder 
    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)
        print(f"Output folder banaya gaya: {OUTPUT_FOLDER}")

    
    try:
        word_app = win32.Dispatch('Word.Application')
        word_app.Visible = False

        df = pd.read_excel(DATA_FILE_PATH, sheet_name=SHEET_NAME)
        
        if df.empty:
            print(f"\nERROR: Excel sheet '{SHEET_NAME}' mein koi data nahi hai.")
            word_app.Quit() 
            sys.exit()

        if FILENAME_COLUMN not in df.columns:
            raise KeyError(f"Configuration mein set kiya gaya column '{FILENAME_COLUMN}' data sheet mein nahi mila.")

        print(f"\n--- Processing Started ---")
        print(f"Loaded {len(df)} records. Saving as PDF...")

        for index, row in df.iterrows():
            
            
            file_name_raw = str(row[FILENAME_COLUMN])
            file_name_clean = clean_filename(file_name_raw)
            temp_docx_path = OUTPUT_FOLDER + file_name_clean + "_TEMP.docx"
            output_pdf_path = OUTPUT_FOLDER + file_name_clean + ".pdf"
            
            
            shutil.copyfile(TEMPLATE_FILE_PATH, temp_docx_path)
            
            
            doc = Document(temp_docx_path) 
            
           
            for key, value in row.items():
                placeholder = f"{{{{{key}}}}}"
                
                # --- DATE FORMATTING LOGIC ---
                replacement = str(value)
              
                if key == 'Date' and pd.notna(value):
                    try:
                        dt_obj = pd.to_datetime(value).to_pydatetime()
                        replacement = dt_obj.strftime('%d %b %Y')
                    except Exception as e:
                        print(f"Warning: Could not format date for column '{key}'. Using original value. Error: {e}")
                # -----------------------------
                
                replace_text_in_all_elements(doc, placeholder, replacement)


            doc.save(temp_docx_path)
            
           
            word_doc = None
            try:
                
                word_doc = word_app.Documents.Open(temp_docx_path)
                
                
                word_doc.ExportAsFixedFormat(
                    OutputFileName=output_pdf_path, 
                    ExportFormat=wdExportFormatPDF
                )
                
                
                word_doc.Close(SaveChanges=0) 
                os.remove(temp_docx_path)
                
                print(f"Saved: {file_name_clean}.pdf")

            except Exception as e:
                print(f"Conversion Error for {file_name_clean}: {e}")
                if word_doc:
                    word_doc.Close(SaveChanges=0)


        word_app.Quit()
        print("\n--- Processing Complete ---")
        print(f"Total {len(df)} files are saved.")

    except Exception as e:
        print(f"\(unexpected) Please check error. Error: {e}")
        if 'word_app' in locals() and word_app:
            word_app.Quit()